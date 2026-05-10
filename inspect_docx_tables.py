from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def inspect_docx(path: Path) -> None:
    doc = Document(path)
    print(f"FILE: {path}")
    print(f"TABLE_COUNT: {len(doc.tables)}")
    for table_index, table in enumerate(doc.tables, start=1):
        row_count = len(table.rows)
        col_count = len(table.columns)
        print(f"TABLE {table_index}: rows={row_count}, cols={col_count}")
        preview_rows = min(3, row_count)
        for row_index in range(preview_rows):
            cell_texts = []
            for col_index, cell in enumerate(table.rows[row_index].cells, start=1):
                normalized = cell.text.replace("\n", "\\n").strip()
                cell_texts.append(f"C{col_index}={normalized!r}")
            print(f"  R{row_index + 1}: " + ", ".join(cell_texts))
    print()


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Print table counts, dimensions, and the first 3 rows of each table in a DOCX file."
    )
    parser.add_argument(
        "paths",
        nargs="*",
        help="One or more .docx files to inspect. Defaults to the requested filename in the current directory.",
    )
    args = parser.parse_args()

    if args.paths:
        paths = [Path(p) for p in args.paths]
    else:
        paths = [Path("刘老师班--运算器实验报告--学生姓名.docx")]

    for path in paths:
        if not path.exists():
            print(f"FILE NOT FOUND: {path}")
            print()
            continue
        inspect_docx(path)


if __name__ == "__main__":
    main()
